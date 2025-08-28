import os, re, json, uuid, asyncio, traceback
from pathlib import Path
from typing import Dict, Any
from datetime import datetime
from html import escape as html_escape

import ffmpeg
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import PlainTextResponse, JSONResponse
from dotenv import load_dotenv
load_dotenv()

# =========================
# ENV (СЕКРЕТЫ ТОЛЬКО ИЗ Render)
# =========================
OPENAI_API_KEY          = os.getenv("OPENAI_API_KEY", "")
TELEGRAM_BOT_TOKEN      = os.getenv("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_WEBHOOK_SECRET = os.getenv("TELEGRAM_WEBHOOK_SECRET", "")
PUBLIC_BASE_URL         = os.getenv("PUBLIC_BASE_URL")  # опционально
RENDER_EXTERNAL_URL     = os.getenv("RENDER_EXTERNAL_URL")  # Render подставит сам
MIN_TEXT_CHARS          = int(os.getenv("MIN_TEXT_CHARS", "120"))
FSM_TIMEOUT             = int(os.getenv("FSM_TIMEOUT", "180"))  # сек
OPENAI_MODEL_CHAT       = os.getenv("OPENAI_MODEL_CHAT", "gpt-4o-mini")
OPENAI_MODEL_TRANSCRIBE = os.getenv("OPENAI_MODEL_TRANSCRIBE", "gpt-4o-transcribe")
OPENAI_MODEL_FALLBACK   = "whisper-1"

# лимиты загрузок
MAX_TG_DOWNLOAD_MB      = int(os.getenv("MAX_TG_DOWNLOAD_MB", "19"))    # лимит Telegram ~20MB
MAX_HTTP_DOWNLOAD_MB    = int(os.getenv("MAX_HTTP_DOWNLOAD_MB", "200")) # по ссылке

if not OPENAI_API_KEY or not TELEGRAM_BOT_TOKEN or not TELEGRAM_WEBHOOK_SECRET:
    raise RuntimeError("Env vars required: OPENAI_API_KEY, TELEGRAM_BOT_TOKEN, TELEGRAM_WEBHOOK_SECRET")

# =====================
# OpenAI async client
# =====================
from openai import AsyncOpenAI
client = AsyncOpenAI(api_key=OPENAI_API_KEY)

QUALITY_SYSTEM_PROMPT = (
    "Ты — строгий контролёр качества разговоров (менеджер ↔ клиент). "
    "Твоя задача: (1) оценить общий уровень, (2) явно определить, было ли возражение о времени/длительности/темпе "
    "и (3) оценить отработку именно этого возражения, (4) выдать рекомендации.\n"
    "Верни СТРОГО ВАЛИДНЫЙ JSON без лишнего текста по схеме:\n"
    "{\n"
    '  "scores": {\n'
    '    "greeting": 0-5,\n'
    '    "needs_analysis": 0-5,\n'
    '    "value_presentation": 0-5,\n'
    '    "objection_time_detected": 0-1,\n'
    '    "objection_time_handling": 0-5,\n'
    '    "next_steps": 0-5,\n'
    '    "empathy_tone": 0-5,\n'
    '    "structure": 0-5,\n'
    '    "overall": 0-100\n'
    "  },\n"
    '  "findings": {\n'
    '    "time_objection_quotes": ["..."],\n'
    '    "key_strengths": ["..."],\n'
    '    "key_issues": ["..."]\n'
    "  },\n"
    '  "actions": ["..."],\n'
    '  "verdict": "1–2 предложения"\n'
    "}\n"
    "Если данных для оценки нет, ставь 0 и укажи причину в key_issues. Пиши только JSON."
)

# ===========
# Файловая СУ
# ===========
DATA_ROOT = Path("./data")
INBOX  = DATA_ROOT / "inbox"
OUTBOX = DATA_ROOT / "outbox"
for p in (DATA_ROOT, INBOX, OUTBOX):
    p.mkdir(parents=True, exist_ok=True)

HISTORY_FILE = DATA_ROOT / "history.json"
if not HISTORY_FILE.exists():
    HISTORY_FILE.write_text("[]", encoding="utf-8")

# ===========
# Утилиты IO
# ===========
AUDIO_EXT = {".mp3",".wav",".m4a",".aac",".ogg",".oga",".flac",".webm"}
VIDEO_EXT = {".mp4",".mov",".mkv",".avi",".webm"}
TEXT_EXT  = {".txt",".md",".csv",".rtf"}

def safe_filename(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_.\-]+", "_", name).strip("._") or f"file_{uuid.uuid4().hex[:6]}"

def is_audio(name: str) -> bool: return Path(name).suffix.lower() in AUDIO_EXT
def is_video(name: str) -> bool: return Path(name).suffix.lower() in VIDEO_EXT
def is_text(name: str)  -> bool: return Path(name).suffix.lower() in TEXT_EXT

def read_text_file(path: Path, limit_mb=15.0) -> str:
    data = path.read_bytes()[:int(limit_mb*1024*1024)]
    return data.decode("utf-8", errors="ignore")

def transcode_to_wav(src_path: Path, sr: int = 16000) -> Path:
    dst_path = src_path.with_suffix(".wav")
    (
        ffmpeg
          .input(str(src_path))
          .output(str(dst_path), ac=1, ar=sr, format="wav")
          .overwrite_output()
          .run(capture_stdout=True, capture_stderr=True)
    )
    return dst_path

def esc(x: object) -> str:
    return html_escape(str(x), quote=True)

def file_preview(path: Path) -> Dict[str, Any]:
    """Короткое превью: имя, размер, тип, длительность (если медиа) или длина текста."""
    info = {"name": path.name, "size_kb": round(path.stat().st_size / 1024, 1)}
    if is_text(path.name):
        content = read_text_file(path, limit_mb=1.5)
        info["type"] = "text"
        info["chars"] = len(content)
        info["snippet"] = (content[:300] + "…") if len(content) > 300 else content
    elif is_audio(path.name) or is_video(path.name):
        try:
            probe = ffmpeg.probe(str(path))
            dur = float(probe["format"].get("duration", 0.0))
            info["type"] = "audio" if is_audio(path.name) else "video"
            info["duration_sec"] = int(dur)
        except Exception:
            info["type"] = "media"
            info["duration_sec"] = None
    else:
        info["type"] = "unknown"
    return info

def add_history_entry(report_id: str, src_name: str, analysis: dict):
    try:
        hist = json.loads(HISTORY_FILE.read_text(encoding="utf-8"))
    except Exception:
        hist = []
    overall = None
    try:
        overall = int(analysis.get("scores", {}).get("overall"))
    except Exception:
        pass
    verdict = analysis.get("verdict", "")
    hist.append({
        "report_id": report_id,
        "src_name": src_name,
        "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "overall": overall,
        "verdict": verdict[:200]
    })
    HISTORY_FILE.write_text(json.dumps(hist, ensure_ascii=False, indent=2), encoding="utf-8")

# формат размера
def fmt_size(bytes_n: int) -> str:
    v = float(bytes_n)
    for unit in ["B","KB","MB","GB"]:
        if v < 1024 or unit == "GB":
            return f"{v:.1f} {unit}"
        v /= 1024

# === HTTP загрузка больших файлов ===
import aiohttp
from urllib.parse import urlparse, unquote

async def download_via_http(url: str, dst: Path, max_mb: int = MAX_HTTP_DOWNLOAD_MB) -> None:
    """Скачивает файл по HTTP(S) с ограничением размера."""
    limit_bytes = max_mb * 1024 * 1024
    dst.parent.mkdir(parents=True, exist_ok=True)
    async with aiohttp.ClientSession() as session:
        async with session.get(url, allow_redirects=True, timeout=aiohttp.ClientTimeout(total=600)) as resp:
            resp.raise_for_status()
            size = 0
            with open(dst, "wb") as f:
                async for chunk in resp.content.iter_chunked(1 << 14):
                    size += len(chunk)
                    if size > limit_bytes:
                        f.close()
                        dst.unlink(missing_ok=True)
                        raise RuntimeError(f"Файл больше лимита {max_mb} МБ")
                    f.write(chunk)

def filename_from_url(url: str, default: str = "file.bin") -> str:
    path = urlparse(url).path
    name = unquote(Path(path).name or default)
    return safe_filename(name)

# =====================
# OpenAI helpers (async)
# =====================
async def openai_transcribe(path: Path) -> str:
    """Транскрипция: сначала prefer модель, затем fallback."""
    try:
        with open(path, "rb") as f:
            r = await client.audio.transcriptions.create(model=OPENAI_MODEL_TRANSCRIBE, file=f)
        text = getattr(r, "text", None) or (r.get("text") if isinstance(r, dict) else None)
        if text: return text.strip()
    except Exception:
        pass
    # fallback
    with open(path, "rb") as f:
        r = await client.audio.transcriptions.create(model=OPENAI_MODEL_FALLBACK, file=f)
    text = getattr(r, "text", None) or (r.get("text") if isinstance(r, dict) else "")
    return (text or "").strip()

async def openai_quality_json(transcript_text: str) -> dict:
    msgs = [
        {"role": "system", "content": QUALITY_SYSTEM_PROMPT},
        {"role": "user",   "content": transcript_text[:20000]}
    ]
    resp = await client.chat.completions.create(
        model=OPENAI_MODEL_CHAT,
        messages=msgs,
        temperature=0.0
    )
    content = resp.choices[0].message.content
    m = re.search(r"\{[\s\S]*\}", content or "")
    raw = m.group(0) if m else (content or "{}")
    try:
        return json.loads(raw)
    except Exception:
        try:
            return json.loads(raw.replace("'", '"'))
        except Exception as e:
            return {"raw": content, "error": f"JSON parse failed: {e}"}

# =====================
# Генераторы отчётов
# =====================
def md_report(analysis: dict, src_name: str) -> str:
    ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    def get(path, default=None):
        cur = analysis
        for k in path:
            if isinstance(cur, dict) and k in cur: cur = cur[k]
            else: return default
        return cur
    lines = [f"# Отчёт по контролю качества ({src_name})", f"_Создано: {ts}_", ""]
    scores = get(["scores"], {})
    if isinstance(scores, dict):
        lines.append("## Оценки")
        for k,v in scores.items(): lines.append(f"- **{k}**: {v}")
        lines.append("")
    findings = get(["findings"], {})
    if isinstance(findings, dict):
        to = findings.get("time_objection_quotes") or []
        if to:
            lines.append("## Фрагменты про возражение времени")
            for q in to[:8]: lines.append(f"> {q}")
            lines.append("")
        ks = findings.get("key_strengths") or []
        if ks:
            lines.append("## Сильные стороны")
            for it in ks[:8]: lines.append(f"- {it}")
            lines.append("")
        ki = findings.get("key_issues") or []
        if ki:
            lines.append("## Проблемы")
            for it in ki[:10]: lines.append(f"- {it}")
            lines.append("")
    actions = get(["actions"], [])
    if actions:
        lines.append("## Рекомендации менеджеру")
        for a in actions[:10]: lines.append(f"- {a}")
        lines.append("")
    verdict = get(["verdict"], "")
    if verdict:
        lines.append("## Вывод"); lines.append(verdict); lines.append("")
    if "error" in analysis:
        lines.append("> ⚠️ JSON не распарсился — приложен сырой ответ в конце.")
        lines.append("")
    return "\n".join(lines)

# DOCX
from docx import Document
def docx_report(analysis: dict, src_name: str, out_path: Path):
    doc = Document()
    doc.add_heading(f"Отчёт по контролю качества ({src_name})", 0)
    ts = datetime.utcnow().strftime("%Y-%m-%d %H:%М UTC")
    doc.add_paragraph(f"Создано: {ts}")

    scores = analysis.get("scores", {})
    if isinstance(scores, dict):
        doc.add_heading("Оценки", level=1)
        for k, v in scores.items():
            p = doc.add_paragraph()
            p.add_run(f"{k}: ").bold = True
            p.add_run(str(v))

    findings = analysis.get("findings", {})
    if isinstance(findings, dict):
        to = findings.get("time_objection_quotes") or []
        if to:
            doc.add_heading("Фрагменты про возражение времени", level=1)
            for q in to[:8]:
                doc.add_paragraph(q)

        ks = findings.get("key_strengths") or []
        if ks:
            doc.add_heading("Сильные стороны", level=1)
            for it in ks[:8]:
                doc.add_paragraph(f"• {it}")

        ki = findings.get("key_issues") or []
        if ki:
            doc.add_heading("Проблемы", level=1)
            for it in ki[:10]:
                doc.add_paragraph(f"• {it}")

    actions = analysis.get("actions") or []
    if actions:
        doc.add_heading("Рекомендации менеджеру", level=1)
        for a in actions[:10]:
            doc.add_paragraph(f"• {a}")

    verdict = analysis.get("verdict")
    if verdict:
        doc.add_heading("Вывод", level=1)
        doc.add_paragraph(verdict)

    if "error" in analysis:
        doc.add_paragraph("⚠️ JSON не распарсился — приложен сырой ответ в конце.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

# PDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
def pdf_report(analysis: dict, src_name: str, out_path: Path):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(out_path), pagesize=A4)
    width, height = A4
    margin = 15 * mm
    x = margin
    y = height - margin

    def writeln(text: str, leading=14):
        nonlocal y
        max_width = width - 2*margin
        from reportlab.pdfbase.pdfmetrics import stringWidth
        words = text.split()
        line = ""
        for w in words:
            cand = (line + " " + w).strip()
            if stringWidth(cand, "Helvetica", 11) <= max_width:
                line = cand
            else:
                c.setFont("Helvetica", 11); c.drawString(x, y, line); y -= leading
                if y < margin: c.showPage(); y = height - margin
                line = w
        if line:
            c.setFont("Helvetica", 11); c.drawString(x, y, line); y -= leading
            if y < margin: c.showPage(); y = height - margin

    ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    c.setFont("Helvetica-Bold", 14); c.drawString(x, y, f"Отчёт по контролю качества ({src_name})"); y -= 18
    c.setFont("Helvetica", 10); c.drawString(x, y, f"Создано: {ts}"); y -= 16

    scores = analysis.get("scores", {})
    if isinstance(scores, dict) and scores:
        c.setFont("Helvetica-Bold", 12); c.drawString(x, y, "Оценки"); y -= 16
        for k, v in scores.items(): writeln(f"- {k}: {v}")

    findings = analysis.get("findings", {})
    if isinstance(findings, dict):
        to = findings.get("time_objection_quotes") or []
        if to:
            c.setFont("Helvetica-Bold", 12); c.drawString(x, y, "Фрагменты про возражение времени"); y -= 16
            for q in to[:8]: writeln(f"» {q}")

        ks = findings.get("key_strengths") or []
        if ks:
            c.setFont("Helvetica-Bold", 12); c.drawString(x, y, "Сильные стороны"); y -= 16
            for it in ks[:8]: writeln(f"• {it}")

        ki = findings.get("key_issues") or []
        if ki:
            c.setFont("Helvetica-Bold", 12); c.drawString(x, y, "Проблемы"); y -= 16
            for it in ki[:10]: writeln(f"• {it}")

    actions = analysis.get("actions") or []
    if actions:
        c.setFont("Helvetica-Bold", 12); c.drawString(x, y, "Рекомендации менеджеру"); y -= 16
        for a in actions[:10]: writeln(f"- {a}")

    verdict = analysis.get("verdict")
    if verdict:
        c.setFont("Helvetica-Bold", 12); c.drawString(x, y, "Вывод"); y -= 16
        writeln(verdict)

    if "error" in analysis:
        writeln("⚠️ JSON не распарсился — приложен сырой ответ в конце.")

    c.showPage()
    c.save()

# ===========================================
# Высокоуровневые обработчики анализа
# ===========================================
from aiogram.types import Message, BufferedInputFile
async def handle_text_and_reply(text: str, message: Message, src_name: str):
    await message.answer("🧠 Анализирую качество...")
    analysis = await openai_quality_json(text)

    out_base = f"report_{uuid.uuid4().hex[:8]}"
    md_path  = OUTBOX / f"{out_base}.md"
    docx_path= OUTBOX / f"{out_base}.docx"
    pdf_path = OUTBOX / f"{out_base}.pdf"

    # Markdown
    md_path.write_text(md_report(analysis, src_name=src_name), encoding="utf-8")
    # DOCX/PDF
    try: docx_report(analysis, src_name=src_name, out_path=docx_path)
    except Exception: docx_path = None
    try: pdf_report(analysis, src_name=src_name, out_path=pdf_path)
    except Exception: pdf_path = None

    verdict = analysis.get("verdict") or "Готов отчёт."
    await message.answer(f"✅ Готово: {verdict}")

    # Отправляем файлы
    with open(md_path, "rb") as f:
        await message.answer_document(BufferedInputFile(f.read(), filename=md_path.name), caption="Отчёт (Markdown)")
    if docx_path and docx_path.exists():
        with open(docx_path, "rb") as f:
            await message.answer_document(BufferedInputFile(f.read(), filename=docx_path.name), caption="Отчёт (DOCX)")
    if pdf_path and pdf_path.exists():
        with open(pdf_path, "rb") as f:
            await message.answer_document(BufferedInputFile(f.read(), filename=pdf_path.name), caption="Отчёт (PDF)")

    # История
    add_history_entry(out_base, src_name, analysis)

async def process_media_file(path: Path, message: Message, orig_name: str):
    await message.answer("🎧 Конвертирую и транскрибирую...")
    try:
        wav = transcode_to_wav(path)
        text = await openai_transcribe(wav)
    except Exception as e:
        await message.answer(f"⚠️ Ошибка транскрипции: {e}")
        return
    if not text.strip():
        await message.answer("⚠️ Не удалось получить текст.")
        return
    await handle_text_and_reply(text, message, src_name=orig_name)

# ===========================================
# Aiogram v3 + FastAPI (guided flow)
# ===========================================
from aiogram import Bot, Dispatcher, F, Router
from aiogram.types import Update, CallbackQuery, InlineKeyboardButton, InlineKeyboardMarkup
from aiogram.filters import CommandStart, Command
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext

bot = Bot(TELEGRAM_BOT_TOKEN)
dp  = Dispatcher(storage=MemoryStorage())
rt  = Router()
dp.include_router(rt)

HELP = (
    "Я — Нейро-контролёр качества.\n"
    "1) Выберите формат: текст или файл\n"
    "2) Пришлите материал\n"
    "3) Подтвердите анализ — пришлю отчёты (MD/DOCX/PDF)\n\n"
    "Команды: /start, /help, /cancel"
)

class Flow(StatesGroup):
    menu = State()
    waiting_text = State()
    waiting_file = State()
    confirm = State()

def kb_main_menu():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🔤 Анализировать текст", callback_data="flow:text")],
        [InlineKeyboardButton(text="📁 Загрузить файл (аудио/видео/док)", callback_data="flow:file")],
    ])

def kb_confirm():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Начать анализ", callback_data="flow:go")],
        [InlineKeyboardButton(text="↩️ Отмена", callback_data="flow:cancel")]
    ])

def kb_cancel():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="↩️ Отмена", callback_data="flow:cancel")]
    ])

async def auto_clear(state: FSMContext, timeout=FSM_TIMEOUT):
    await asyncio.sleep(timeout)
    if await state.get_state():
        await state.clear()

@rt.message(CommandStart())
async def cmd_start(msg: Message, state: FSMContext):
    await state.clear()
    await state.set_state(Flow.menu)
    asyncio.create_task(auto_clear(state))
    await msg.answer("Привет! Выберите вариант:", reply_markup=kb_main_menu())

@rt.message(Command("help"))
async def cmd_help(msg: Message, state: FSMContext):
    await msg.answer(HELP, reply_markup=kb_main_menu())

@rt.message(Command("cancel"))
async def cmd_cancel(msg: Message, state: FSMContext):
    await state.clear()
    await msg.answer("Отмена. Вернулись в меню.", reply_markup=kb_main_menu())

@rt.callback_query(F.data == "flow:text")
async def cb_text(cb: CallbackQuery, state: FSMContext):
    await state.set_state(Flow.waiting_text)
    asyncio.create_task(auto_clear(state))
    await cb.message.edit_text(
        f"Пришлите текст (минимум {MIN_TEXT_CHARS} символов).",
        reply_markup=kb_cancel()
    )
    await cb.answer()

@rt.callback_query(F.data == "flow:file")
async def cb_file(cb: CallbackQuery, state: FSMContext):
    await state.set_state(Flow.waiting_file)
    asyncio.create_task(auto_clear(state))
    await cb.message.edit_text(
        "Пришлите файл: 🎧 аудио / 🎥 видео / 📄 документ (txt/md/csv/rtf).",
        reply_markup=kb_cancel()
    )
    await cb.answer()

@rt.callback_query(F.data == "flow:cancel")
async def cb_cancel(cb: CallbackQuery, state: FSMContext):
    await state.clear()
    await cb.message.edit_text("Отменено. Меню:", reply_markup=kb_main_menu())
    await cb.answer()

URL_RE = re.compile(r"https?://[^\s]+", re.IGNORECASE)

@rt.message(Flow.waiting_text, F.text)
async def handle_text_intake(msg: Message, state: FSMContext):
    txt = (msg.text or "").strip()
    if len(txt) < MIN_TEXT_CHARS:
        await msg.answer(
            f"Текста маловато ({len(txt)}). Минимум {MIN_TEXT_CHARS}.",
            reply_markup=kb_cancel()
        )
        return
    await state.update_data(src_type="text", text=txt, src_name="text.txt")
    await state.set_state(Flow.confirm)
    preview_html = f"<b>Текст:</b> {len(txt)} символов\n\nНажмите «Начать анализ»."
    await msg.answer(preview_html, reply_markup=kb_confirm(), parse_mode="HTML")

@rt.message(Flow.waiting_file, F.document | F.audio | F.voice | F.video | F.video_note)
async def handle_file_intake(msg: Message, state: FSMContext):
    try:
        file = msg.document or msg.audio or msg.voice or msg.video or msg.video_note
        file_size = getattr(file, "file_size", None) or getattr(file, "length", None) or 0
        if file_size and file_size > MAX_TG_DOWNLOAD_MB * 1024 * 1024:
            human = fmt_size(file_size)
            await state.update_data(expect_url=True)
            await msg.answer(
                (
                    f"Файл слишком большой для Bot API ({human} > {MAX_TG_DOWNLOAD_MB} MB).\n\n"
                    "Пришлите <b>ссылку</b> на файл (Google Drive / Яндекс.Диск / Dropbox / прямой URL), "
                    f"я скачаю его напрямую (до {MAX_HTTP_DOWNLOAD_MB} MB)."
                ),
                parse_mode="HTML",
                reply_markup=kb_cancel()
            )
            return

        raw_name = getattr(file, "file_name", None) or "file.bin"
        fname = safe_filename(raw_name)
        dst = INBOX / fname
        await bot.download(file, destination=dst)

        info = file_preview(dst)
        lines = [
            f"<b>Файл:</b> {esc(info.get('name'))} ({esc(info.get('size_kb'))} KB)",
            f"<b>Тип:</b> {esc(info.get('type'))}",
        ]
        if info.get("duration_sec"):
            lines.append(f"<b>Длительность:</b> {esc(info['duration_sec'])} сек")
        if info.get("type") == "text" and info.get("chars") is not None:
            lines.append(f"<b>Длина:</b> {esc(info['chars'])} символов")
        preview_html = "\n".join(lines) + "\n\nНажмите «Начать анализ»."

        await state.update_data(src_type="file", file_path=str(dst), src_name=fname, expect_url=False)
        await state.set_state(Flow.confirm)
        await msg.answer(preview_html, reply_markup=kb_confirm(), parse_mode="HTML")

    except Exception as e:
        await msg.answer(f"Ошибка при получении файла: {e}", reply_markup=kb_cancel())

@rt.message(Flow.waiting_file, F.text)
async def handle_url_instead_of_file(msg: Message, state: FSMContext):
    data = await state.get_data()
    if not data.get("expect_url"):
        await msg.answer("Жду файл (или нажмите Отмена).", reply_markup=kb_cancel())
        return

    txt = (msg.text or "").strip()
    m = URL_RE.search(txt)
    if not m:
        await msg.answer("Пришлите, пожалуйста, прямую ссылку на файл.", reply_markup=kb_cancel())
        return

    url = m.group(0)
    try:
        fname = filename_from_url(url)
        dst = INBOX / fname
        await msg.answer("⬇️ Скачиваю файл по ссылке, подождите…")
        await download_via_http(url, dst, max_mb=MAX_HTTP_DOWNLOAD_MB)

        info = file_preview(dst)
        lines = [
            f"<b>Файл:</b> {esc(info.get('name'))} ({esc(info.get('size_kb'))} KB)",
            f"<b>Тип:</b> {esc(info.get('type'))}",
        ]
        if info.get("duration_sec"):
            lines.append(f"<b>Длительность:</b> {esc(info['duration_sec'])} сек")
        if info.get("type") == "text" and info.get("chars") is not None:
            lines.append(f"<b>Длина:</b> {esc(info['chars'])} символов")
        preview_html = "\n".join(lines) + "\n\nНажмите «Начать анализ»."

        await state.update_data(src_type="file", file_path=str(dst), src_name=fname, expect_url=False)
        await state.set_state(Flow.confirm)
        await msg.answer(preview_html, reply_markup=kb_confirm(), parse_mode="HTML")

    except Exception as e:
        await msg.answer(f"Не удалось скачать файл по ссылке: {e}", reply_markup=kb_cancel())

@rt.callback_query(Flow.confirm, F.data == "flow:go")
async def cb_go_analyze(cb: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    src_type = data.get("src_type")
    src_name = data.get("src_name") or "source"
    await cb.message.edit_text("🧠 Анализирую, подождите...")

    try:
        if src_type == "text":
            text = data.get("text", "")
            await handle_text_and_reply(text, cb.message, src_name=src_name)
        elif src_type == "file":
            path = Path(data.get("file_path"))
            if is_text(path.name):
                t = read_text_file(path)
                await handle_text_and_reply(t, cb.message, src_name=src_name)
            else:
                await process_media_file(path, cb.message, orig_name=src_name)
        else:
            await cb.message.answer("Не понял тип источника. Начните заново.", reply_markup=kb_main_menu())
        await state.clear()
        await state.set_state(Flow.menu)
    except Exception as e:
        traceback.print_exc()
        await cb.message.answer(f"Ошибка при анализе: {e}", reply_markup=kb_main_menu())
        await state.clear()
        await state.set_state(Flow.menu)

    await cb.answer()

@rt.message(F.text)
async def fallback_text(msg: Message, state: FSMContext):
    if not await state.get_state():
        await state.set_state(Flow.menu)
        await msg.answer("Выберите вариант:", reply_markup=kb_main_menu())

# ===========================
# FastAPI + webhook для Render
# ===========================
app = FastAPI()

@app.get("/", response_class=PlainTextResponse)
async def root():
    return "OK: neuro-qc-bot webhook"

def build_webhook_url() -> str:
    base = PUBLIC_BASE_URL or RENDER_EXTERNAL_URL
    if not base:
        raise RuntimeError("No PUBLIC_BASE_URL or RENDER_EXTERNAL_URL provided by Render.")
    return base.rstrip("/") + f"/webhook/{TELEGRAM_WEBHOOK_SECRET}"

@app.on_event("startup")
async def on_startup():
    url = build_webhook_url()
    await bot.set_webhook(
        url,
        secret_token=TELEGRAM_WEBHOOK_SECRET,
        drop_pending_updates=True,
        allowed_updates=["message","edited_message","callback_query"]
    )
    print("Webhook set:", url)

from aiogram.types import Update
@app.post("/webhook/{secret}")
async def telegram_webhook(secret: str, request: Request):
    if secret != TELEGRAM_WEBHOOK_SECRET:
        raise HTTPException(status_code=403, detail="forbidden")
    data = await request.json()
    update = Update.model_validate(data)
    await dp.feed_update(bot, update)
    return JSONResponse({"ok": True})
